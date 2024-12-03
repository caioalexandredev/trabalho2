import requests
from decouple import config
from docx import Document
from io import BytesIO

class GalileuAPIService:
    def __init__(self):
        self.api_host = config('API_GALILEU_HOST').rstrip('/')
        self.api_bearer_token = config('API_GALILEU_BEARER_TOKEN')
        self.headers = {
            'Authorization': f'Bearer {self.api_bearer_token}',
        }

    def get_procedimento_pericial(self, procedimento_id):
        endpoint = f'procedimentopericial/{procedimento_id}'
        url = f'{self.api_host}/{endpoint}'

        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            return response.json()
        except requests.RequestException as e:
            return {'error': str(e)}


class WordFileGenerator:
    @staticmethod
    def generate_word_file_test(data):
        doc = Document()
        
        doc.add_heading(data.get('title', 'Relatório'), level=1)
        
        for section_title, section_content in data.get('sections', {}).items():
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(section_content)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer